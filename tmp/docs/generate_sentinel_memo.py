from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(r"C:\springcloud\shop\shop")
DOCX_PATH = ROOT / "output" / "doc" / "sentinel-gateway-config-memo.docx"
SVG_PATH = ROOT / "gateway" / "docs" / "sentinel-gateway-rule-map.svg"


APPLICATION_YAML = """server:
  port: 8080
spring:
  application:
    name: gateway
  config:
    import:
      - nacos:sentinel?group=commons
  cloud:
    sentinel:
      filter:
        enabled: false
      datasource:
        gatewayGwFlow:
          nacos:
            server-addr: ${spring.cloud.nacos.server-addr}
            namespace: ${spring.cloud.nacos.config.namespace}
            data-id: gateway-gw-flow-rules
            group-id: app
            data-type: json
            rule-type: gw-flow
        gatewayApiGroup:
          nacos:
            server-addr: ${spring.cloud.nacos.server-addr}
            namespace: ${spring.cloud.nacos.config.namespace}
            data-id: gateway-gw-api-group-rules
            group-id: app
            data-type: json
            rule-type: gw-api-group"""


GATEWAY_CONFIG_JAVA = """@Configuration
public class GatewayConfig {

    @PostConstruct
    public void initBlockHandler() {
        GatewayCallbackManager.setBlockHandler(urlBlockHandler());
    }

    public BlockRequestHandler urlBlockHandler() {
        return (exchange, e) -> {
            String msg = buildMessage(e);
            String path = exchange.getRequest().getURI().getPath();
            String result = String.format(
                    "{\\"code\\":429,\\"msg\\":\\"%s\\",\\"data\\":\\"%s\\"}",
                    msg,
                    path
            );
            return ServerResponse.status(HttpStatus.TOO_MANY_REQUESTS)
                    .contentType(MediaType.APPLICATION_JSON)
                    .body(BodyInserters.fromValue(result));
        };
    }
}"""


API_GROUP_JSON = """[
  {
    "apiName": "auth-api",
    "predicateItems": [
      {
        "pattern": "/auth/**",
        "matchStrategy": 1
      }
    ]
  },
  {
    "apiName": "order-api",
    "predicateItems": [
      {
        "pattern": "/order/**",
        "matchStrategy": 1
      }
    ]
  },
  {
    "apiName": "flash-sale-api",
    "predicateItems": [
      {
        "pattern": "/flash-sale/**",
        "matchStrategy": 1
      }
    ]
  },
  {
    "apiName": "cart-api",
    "predicateItems": [
      {
        "pattern": "/cart/**",
        "matchStrategy": 1
      }
    ]
  }
]"""


GW_FLOW_JSON = """[
  {
    "resource": "auth",
    "resourceMode": 0,
    "grade": 1,
    "count": 20,
    "intervalSec": 1,
    "controlBehavior": 0
  },
  {
    "resource": "order",
    "resourceMode": 0,
    "grade": 1,
    "count": 10,
    "intervalSec": 1,
    "controlBehavior": 0
  },
  {
    "resource": "flash-sale",
    "resourceMode": 0,
    "grade": 1,
    "count": 5,
    "intervalSec": 1,
    "controlBehavior": 0
  },
  {
    "resource": "auth-api",
    "resourceMode": 1,
    "grade": 1,
    "count": 15,
    "intervalSec": 1,
    "controlBehavior": 0
  },
  {
    "resource": "order-api",
    "resourceMode": 1,
    "grade": 1,
    "count": 8,
    "intervalSec": 1,
    "controlBehavior": 0
  },
  {
    "resource": "flash-sale-api",
    "resourceMode": 1,
    "grade": 1,
    "count": 3,
    "intervalSec": 1,
    "controlBehavior": 0
  },
  {
    "resource": "cart-api",
    "resourceMode": 1,
    "grade": 1,
    "count": 1,
    "intervalSec": 1,
    "controlBehavior": 0
  }
]"""


ROUTES = [
    ("ruoyi", "/ruoyi/**"),
    ("product", "/product/**"),
    ("file", "/file/**"),
    ("member", "/member/**"),
    ("ware", "/ware/**"),
    ("search", "/search/**"),
    ("auth", "/auth/**"),
    ("cart", "/cart/**"),
    ("order", "/order/**"),
    ("coupon", "/coupon/**"),
    ("flash-sale", "/flash-sale/**"),
]


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, size=10.5, color="24324A"):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    p.style = doc.styles[f"Heading {level}"]
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    return p


def add_body(doc, text, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10.5)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10.5)
    return p


def add_code_block(doc, title, code):
    add_body(doc, title, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(8)
    for idx, line in enumerate(code.splitlines()):
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(9.5)
        if idx < len(code.splitlines()) - 1:
            run.add_break()


def build_doc():
    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Sentinel 网关限流配置备忘录")
    run.bold = True
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(20)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("适用模块：gateway | 目标：把当前项目的 Sentinel Gateway 限流链路说清楚")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor.from_string("5C6B82")

    add_heading(doc, "1. 一页结论", 1)
    add_bullet(doc, "当前项目的 Sentinel 网关限流分两层：Route ID 做路由级总闸门，API 分组做路径级热点保护。")
    add_bullet(doc, "Nacos 中需要两份规则：gateway-gw-flow-rules 与 gateway-gw-api-group-rules，group 为 app。")
    add_bullet(doc, "application.yaml 中 rule-type 必须分别是 gw-flow 与 gw-api-group，不能写成普通 flow。")
    add_bullet(doc, "被限流后的返回由 GatewayConfig 中的 GatewayCallbackManager.setBlockHandler(...) 统一处理。")

    add_heading(doc, "2. 当前项目最小可用配置链路", 1)
    add_body(doc, "按下面顺序检查，任意一步缺失都会导致“配置没用”或规则不生效。")

    chain = doc.add_table(rows=1, cols=3)
    chain.alignment = WD_TABLE_ALIGNMENT.CENTER
    chain.style = "Table Grid"
    headers = ["环节", "当前要求", "备注"]
    for cell, text in zip(chain.rows[0].cells, headers):
        shade_cell(cell, "EAF2FB")
        set_cell_text(cell, text, bold=True)
    rows = [
        ("依赖", "gateway/pom.xml 包含 sentinel-gateway 与 sentinel-datasource-nacos", "缺少 datasource-nacos 会直接启动报错"),
        ("配置", "application.yaml 中配置 gw-flow 与 gw-api-group 数据源", "Nacos namespace 复用 dev"),
        ("规则来源", "Nacos dataId：gateway-gw-flow-rules / gateway-gw-api-group-rules", "group 使用 app"),
        ("回调", "GatewayConfig 注册 GatewayCallbackManager.setBlockHandler", "统一返回 429 JSON"),
        ("控制台", "API 管理定义分组，网关流控规则绑定 Route ID 或 API 分组", "同一请求可同时命中两类规则"),
    ]
    for row in rows:
        cells = chain.add_row().cells
        for cell, value in zip(cells, row):
            set_cell_text(cell, value)

    add_heading(doc, "3. 当前网关路由与推荐关注点", 1)
    route_table = doc.add_table(rows=1, cols=3)
    route_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    route_table.style = "Table Grid"
    route_headers = ["Route ID", "Path", "建议"]
    for cell, text in zip(route_table.rows[0].cells, route_headers):
        shade_cell(cell, "EAF2FB")
        set_cell_text(cell, text, bold=True)
    for route_id, path in ROUTES:
        cells = route_table.add_row().cells
        recommendation = "重点保护" if route_id in {"auth", "order", "flash-sale", "cart"} else "常规路由"
        for cell, value in zip(cells, (route_id, path, recommendation)):
            set_cell_text(cell, value)

    add_heading(doc, "4. application.yaml 关键配置", 1)
    add_code_block(doc, "核心配置片段", APPLICATION_YAML)

    add_heading(doc, "5. GatewayConfig 关键逻辑", 1)
    add_code_block(doc, "限流回调注册片段", GATEWAY_CONFIG_JAVA)
    add_bullet(doc, "这里只负责限流后的返回内容，不负责写死规则。")
    add_bullet(doc, "规则本体由 Nacos 数据源加载，避免每次改阈值都重启网关。")

    add_heading(doc, "6. Nacos 规则约定", 1)
    rules_table = doc.add_table(rows=1, cols=4)
    rules_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    rules_table.style = "Table Grid"
    for cell, text in zip(rules_table.rows[0].cells, ["dataId", "group", "rule-type", "用途"]):
        shade_cell(cell, "EAF2FB")
        set_cell_text(cell, text, bold=True)
    rules = [
        ("gateway-gw-api-group-rules", "app", "gw-api-group", "定义 auth-api、order-api、flash-sale-api、cart-api 等路径分组"),
        ("gateway-gw-flow-rules", "app", "gw-flow", "定义 Route ID 与 API 分组的限流阈值"),
    ]
    for row in rules:
        cells = rules_table.add_row().cells
        for cell, value in zip(cells, row):
            set_cell_text(cell, value)

    add_code_block(doc, "API 分组示例 JSON", API_GROUP_JSON)
    add_code_block(doc, "网关流控规则示例 JSON", GW_FLOW_JSON)

    add_heading(doc, "7. 控制台怎么看当前规则", 1)
    explain = doc.add_table(rows=1, cols=4)
    explain.alignment = WD_TABLE_ALIGNMENT.CENTER
    explain.style = "Table Grid"
    for cell, text in zip(explain.rows[0].cells, ["请求示例", "命中 Route ID", "命中 API 分组", "通常谁先拦截"]):
        shade_cell(cell, "EAF2FB")
        set_cell_text(cell, text, bold=True)
    examples = [
        ("/auth/login", "auth", "auth-api", "auth-api，因示例阈值 15 小于 20"),
        ("/order/submit", "order", "order-api", "order-api，因示例阈值 8 小于 10"),
        ("/flash-sale/kill", "flash-sale", "flash-sale-api", "flash-sale-api，因示例阈值 3 小于 5"),
        ("/cart/list", "当前截图未见 cart 路由级规则", "cart-api", "cart-api"),
    ]
    for row in examples:
        cells = explain.add_row().cells
        for cell, value in zip(cells, row):
            set_cell_text(cell, value)
    add_body(doc, f"配套 SVG 总览图位置：{SVG_PATH}")

    add_heading(doc, "8. 常见坑与排障", 1)
    pitfalls = [
        "resource 写成 /auth/** 这类路径时，Route ID 规则不会生效；Route 级 resource 必须写 routeId。",
        "gw-flow 与 gw-api-group 的 rule-type 写错时，控制台和数据源都可能看起来“有配置但没效果”。",
        "缺少 com.alibaba.csp:sentinel-datasource-nacos 依赖时，会出现 NacosDataSource 类找不到的启动错误。",
        "同一请求命中多条规则时，只要任意一条先超过阈值，请求就会被拦截。",
        "spring.cloud.sentinel.filter.enabled 在 Gateway 场景下保持 false，避免按普通 Web 资源埋点。",
    ]
    for item in pitfalls:
        add_bullet(doc, item)

    add_heading(doc, "9. 联调验证步骤", 1)
    steps = [
        "启动 gateway，确认应用已在 Sentinel 控制台出现。",
        "在 API 管理中创建 auth-api、order-api、flash-sale-api、cart-api。",
        "在网关流控规则中分别创建 Route ID 与 API 分组规则。",
        "用压测或并发请求访问 /auth/**、/order/**、/flash-sale/**、/cart/**。",
        "观察 429 返回体是否为 {\"code\":429,\"msg\":\"...\",\"data\":\"请求路径\"}。",
    ]
    for idx, step in enumerate(steps, start=1):
        add_body(doc, f"{idx}. {step}")

    add_heading(doc, "10. 交付说明", 1)
    add_bullet(doc, "本文档为当前项目实战备忘录，不是 Sentinel 全量手册。")
    add_bullet(doc, "当前环境缺少 soffice / pdftoppm，未执行页面渲染预览；如需最终版排版校验，请在本机用 Word 打开检查一次。")
    add_bullet(doc, "如后续新增更多 Route ID 规则，可继续复用本备忘录结构扩展。")

    final_section = doc.add_section(WD_SECTION.NEW_PAGE)
    final_section.top_margin = Cm(2.0)
    final_section.bottom_margin = Cm(1.8)
    final_section.left_margin = Cm(2.0)
    final_section.right_margin = Cm(2.0)
    add_heading(doc, "附录：快速检查清单", 1)
    checklist = [
        "依赖是否包含 sentinel-datasource-nacos",
        "application.yaml 是否配置 gw-flow / gw-api-group",
        "Nacos dataId 是否放在 app 组",
        "Route 级 resource 是否写 routeId",
        "429 JSON 是否已被 GatewayConfig 自定义覆盖",
    ]
    for item in checklist:
        add_bullet(doc, item)

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_doc()
    print(DOCX_PATH)
